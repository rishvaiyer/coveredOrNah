from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parents[1] / "output"
OUT_PATH = OUT_DIR / "Formulary_Finder_Clinician_User_Guide.docx"
ASSET_DIR = OUT_DIR / "guide-assets"

INK = "163A3A"
TEAL = "0B6D69"
MUTED = "557573"
GOLD = "9B6A20"
PALE_TEAL = "EAF6F3"
PALE_GOLD = "FFF8E8"
PALE_ROW = "F8FBFA"
LINE = "C8DFDA"
WHITE = "FFFFFF"


def set_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=105, start=140, bottom=105, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        item = borders.find(tag)
        if item is None:
            item = OxmlElement(f"w:{edge}")
            borders.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), size)
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_text(cell, text, size=10.5, color=INK, bold=False):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    return p


def add_paragraph(doc, text="", size=11, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if text:
        set_font(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    p.paragraph_format.line_spacing = 1.05
    set_font(p.add_run(text), size=16 if level == 1 else 12.5, color=TEAL if level == 1 else INK, bold=True)
    return p


def add_note(doc, title, text, fill=PALE_TEAL, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(title + "  "), size=10.5, color=accent, bold=True)
    set_font(p.add_run(text), size=10.5, color=INK)
    return table


def add_cover_panel(doc):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, TEAL)
    set_cell_margins(cell, top=270, start=300, bottom=265, end=300)
    set_cell_border(cell, color=TEAL, size="4")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run("CLINICIAN QUICK GUIDE"), size=9, color="CDEBE5", bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run("Formulary Finder"), size=28, color=WHITE, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run("Plan-first medication coverage workflows for doctors and nurses"), size=12.5, color="DDF2ED")
    return table


def add_table(doc, headers, rows, widths, label_width=False):
    table = doc.add_table(rows=1, cols=len(headers))
    header = table.rows[0]
    mark_header_row(header)
    for index, text in enumerate(headers):
        set_cell_shading(header.cells[index], PALE_TEAL)
        add_text(header.cells[index], text, size=10, color=TEAL, bold=True)
    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        for index, value in enumerate(row_values):
            if row_index % 2 == 1:
                set_cell_shading(row.cells[index], PALE_ROW)
            add_text(row.cells[index], value, size=10.2, bold=(label_width and index == 0))
    set_table_geometry(table, widths)
    return table


def add_screenshot(doc, filename, caption, crop_bottom=0):
    image_path = ASSET_DIR / filename
    if not image_path.exists():
        raise FileNotFoundError(f"Missing guide screenshot: {image_path}")
    label = add_paragraph(doc, "IN-PORTAL EXAMPLE", size=8.2, color=GOLD, bold=True, after=3)
    label.paragraph_format.keep_with_next = True
    inline_shape = doc.add_picture(str(image_path), width=Inches(6.5))
    if crop_bottom:
        blip_fill = inline_shape._inline.graphic.graphicData.pic.blipFill
        src_rect = OxmlElement("a:srcRect")
        src_rect.set("b", str(int(crop_bottom * 100000)))
        blip_fill.insert(1, src_rect)
        inline_shape.height = Inches(6.5 * (720 * (1 - crop_bottom)) / 1280)
    inline_shape._inline.docPr.set("descr", caption)
    p = add_paragraph(doc, caption, size=9.5, color=MUTED, italic=True, after=7)
    p.paragraph_format.keep_with_next = True


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_furniture(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("FORMULARY FINDER  |  CLINICIAN USER GUIDE"), size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    set_font(footer.add_run("August 2026  |  Page "), size=8.5, color=MUTED)
    add_page_field(footer)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_furniture(doc)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    # Page 1
    add_cover_panel(doc)
    add_paragraph(doc, "A concise visual guide to selecting the right medication-coverage workflow before a pharmacy call.", size=12.2, color=MUTED, after=10)
    add_note(doc, "Purpose", "Start with the plan card, then use exact plan evidence when it is available. The portal keeps the lookup in one clinical workflow instead of sending staff to payer websites.")
    add_heading(doc, "The 60-second workflow")
    add_table(
        doc,
        ["Step", "What to do in Formulary Finder"],
        [
            ("1. Open By plan", "Start with the insurer on the card. Do not start on a payer website."),
            ("2. Enter plan-level details", "Choose coverage type, then type the plan or formulary name. Use the pharmacy-benefit label if it is printed on the card."),
            ("3. Match plan", "The portal gives one result: exact imported formulary, exact Medicare Advantage or standalone Part D plan search, or not imported yet."),
            ("4. Search the medicine", "Confirm the medication product, strength, device, tier, and restrictions before acting."),
        ],
        [1700, 7660],
        label_width=True,
    )
    doc.add_page_break()
    add_heading(doc, "What each result means")
    add_table(
        doc,
        ["Portal result", "How to use it"],
        [
            ("Exact imported formulary", "Continue to medication search. Review source date, tier or preferred status, and PA, step therapy, or quantity-limit flags."),
            ("Exact Medicare plan search", "Choose Medicare Advantage or standalone Part D first, then select the exact plan from the correct card. CMS returns product candidates only. Verify county or service area, enrollment or eligibility, product, device, strength, and NDC context."),
            ("Not imported yet", "Stop. Do not infer coverage from the insurer name, network logo, or another plan. Use the card's pharmacy-benefit instructions or your established clinic process when needed."),
            ("Original or Railroad Medicare card", "Use the patient's separate standalone Part D prescription-drug plan card. Original or Railroad Medicare alone does not identify the outpatient prescription plan."),
        ],
        [2500, 6860],
        label_width=True,
    )
    add_note(doc, "Quick safety note", "The portal is a formulary evidence tool. It does not verify eligibility, member benefits, cost sharing, clinical appropriateness, or a final coverage decision.", fill=PALE_GOLD, accent=GOLD)

    doc.add_page_break()

    # Page 2
    add_paragraph(doc, "VISUAL WALKTHROUGH", size=9.5, color=TEAL, bold=True, after=4)
    add_heading(doc, "Screen 1: Start with the insurance card", level=1)
    add_paragraph(doc, "Open By plan. Type the insurer, choose coverage type, then enter the plan or formulary name printed on the card. The pharmacy-benefit label is optional but useful when shown.", after=8)
    add_screenshot(doc, "01-plan-intake.png", "The plan-intake screen is the default starting point. Autocomplete is available for insurer, plan or formulary name, and pharmacy-benefit label.")
    add_note(doc, "Enter only plan details", "Do not enter member IDs, dates of birth, claim numbers, diagnoses, or patient information.")

    doc.add_page_break()

    # Page 3
    add_paragraph(doc, "VISUAL WALKTHROUGH", size=9.5, color=TEAL, bold=True, after=4)
    add_heading(doc, "Screen 2: When an exact plan is not imported", level=1)
    add_paragraph(doc, "Example: UnitedHealthcare Choice Plus. The portal correctly stops instead of substituting a different UnitedHealthcare or Oxford drug list.", after=8)
    add_screenshot(doc, "02-unconfirmed-plan.png", "Read this as unconfirmed, not a denial. Use the card's pharmacy-benefit route or your established clinic process for the next check.", crop_bottom=0.14)
    add_note(doc, "Do not substitute another plan", "A carrier name, network logo, or static shortcut card is not patient-specific coverage evidence.", fill=PALE_GOLD, accent=GOLD)

    doc.add_page_break()

    # Page 4
    add_paragraph(doc, "VISUAL WALKTHROUGH", size=9.5, color=TEAL, bold=True, after=4)
    add_heading(doc, "Screen 3: Exact Medicare Advantage plan", level=1)
    add_paragraph(doc, "Choose Medicare Advantage, then search the carrier, plan name, or contract-plan ID and select the exact plan returned from the CMS data. Do not use a standalone Part D card in this path.", after=8)
    add_screenshot(doc, "03-cms-plan-selection.png", "The plan name, contract-plan ID, and formulary ID confirm the selection. Then search the medicine and verify candidate product, device, strength, NDC, and restrictions.")
    add_note(doc, "CMS boundary", "The finder is New Jersey-focused and does not verify county availability, enrollment, or eligibility.", fill=PALE_GOLD, accent=GOLD)

    doc.add_page_break()

    # Page 6
    add_paragraph(doc, "VISUAL WALKTHROUGH", size=9.5, color=TEAL, bold=True, after=4)
    add_heading(doc, "Screen 4: Exact standalone Part D plan", level=1)
    add_paragraph(doc, "For Original or Railroad Medicare, use the patient's separate prescription-drug plan card. Choose Standalone Part D, search the carrier, plan name, or contract-plan ID, and select the exact plan returned from the CMS data.", after=8)
    add_screenshot(doc, "04-pdp-plan-selection.png", "The separate Part D card, plan name, contract-plan ID, and formulary ID confirm the selection. Original or Railroad Medicare alone does not identify outpatient prescription coverage.")
    add_note(doc, "Keep the paths distinct", "Do not enter a standalone Part D plan in the Medicare Advantage path. A missing plan or product match is unconfirmed, not a denial.", fill=PALE_GOLD, accent=GOLD)

    doc.add_page_break()

    # Page 7
    add_paragraph(doc, "SPECIAL COVERAGE PATHS", size=9.5, color=TEAL, bold=True, after=4)
    add_heading(doc, "Medicare", level=1)
    add_table(
        doc,
        ["Situation", "Correct clinical workflow"],
        [
            ("Medicare Advantage", "Choose Medicare Advantage in By plan. Search by carrier, plan name, or contract-plan ID, then choose the exact CMS plan shown on the card. Confirm county, enrollment or eligibility, and product details separately."),
            ("Original or Railroad Medicare", "Ask for the separate prescription-drug plan card. Choose Standalone Part D, search by carrier, plan name, or contract-plan ID, and select the exact CMS plan. The red-white-blue Medicare card alone does not identify outpatient prescription coverage."),
            ("Standalone prescription-drug plan", "Choose Standalone Part D in the Medicare finder and use the exact separate Part D card. Keep this path distinct from Medicare Advantage and treat a missing match as unconfirmed, not a denial."),
        ],
        [2700, 6660],
        label_width=True,
    )
    add_heading(doc, "When the card is not a drug benefit", level=1)
    add_table(
        doc,
        ["Card type", "What to do"],
        [
            ("Network or administrator", "Identify the underlying payer and PBM. Network participation is not a formulary match."),
            ("Workers' compensation or MVA", "Confirm the injury authorization and pharmacy channel. Do not enter claim numbers in Formulary Finder."),
            ("VA Community Care", "Confirm authorization and use the VA-directed pharmacy channel. Community Care participation does not prove outpatient drug coverage."),
            ("TRICARE or US Family Health Plan", "Confirm the named plan variant and prescription route. Do not substitute a commercial or Medicare formulary."),
        ],
        [2700, 6660],
        label_width=True,
    )
    add_note(doc, "If the exact plan is unavailable", "Use the result as a stop sign, not a denial. The next step is the plan's documented pharmacy-benefit route or your established clinic process.", fill=PALE_GOLD, accent=GOLD)

    doc.add_page_break()

    # Page 8
    add_paragraph(doc, "SCOPE AND CLINIC CHECK", size=9.5, color=TEAL, bold=True, after=4)
    add_heading(doc, "What is loaded today", level=1)
    add_paragraph(doc, "The live pilot currently includes 85 medication records and 16 New Jersey-focused formulary baselines. The New Jersey-focused CMS finder has separate Medicare Advantage and standalone Part D paths and does not verify county availability, enrollment, or eligibility. The insurer directory identifies network participation only, not medication coverage.", after=8)
    add_table(
        doc,
        ["Named baseline formularies", "Region / plan family"],
        [
            ("Horizon Marketplace; Horizon Classic; UHC Commercial; Oxford Freedom", "NJ commercial and marketplace plan families"),
            ("Aetna HMO; Wellcare NJ; Humana NJ; Braven NJ; HealthSpring NJ; Clover NJ", "Named Medicare plan families; use exact CMS plan selection when available"),
            ("AmeriHealth NJ; Cigna 3-Tier; Oscar NJ; Wellpoint NJ FamilyCare", "Named commercial, individual, and NJ FamilyCare plan families"),
        ],
        [4700, 4660],
        label_width=True,
    )
    add_heading(doc, "Before documenting a coverage call", level=1)
    add_table(
        doc,
        ["Confirm", "Do not assume"],
        [
            ("Exact insurer, coverage type, and plan family", "A carrier name alone proves medication coverage"),
            ("Medication product, strength, device, and NDC context", "A generic-name match is always the correct product"),
            ("Restrictions and source date", "Not listed means not covered"),
            ("Clinic policy before any PA step", "The portal submits prior authorization or replaces clinical judgment"),
        ],
        [4680, 4680],
        label_width=True,
    )
    add_paragraph(doc, "Live portal: formulary-finder-pilot-production.up.railway.app", size=10.5, color=TEAL, bold=True, after=2)
    add_paragraph(doc, "Guide version: August 12, 2026. Evidence sources and plan formularies change. Recheck the source date shown in the portal before acting.", size=9.5, color=MUTED, italic=True, after=0)

    doc.core_properties.title = "Formulary Finder Clinician User Guide"
    doc.core_properties.subject = "Plan-level formulary workflow guide"
    doc.core_properties.author = "Formulary Finder"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
